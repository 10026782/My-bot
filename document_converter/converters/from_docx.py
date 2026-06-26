from __future__ import annotations

from pathlib import Path

from ..result import failure, success
from .pandoc import try_pandoc


def convert(source: Path, output: Path, input_type: str, output_type: str):
    if try_pandoc(source, output, input_type, output_type):
        return success(str(output), ["converted with pandoc"])

    from docx import Document

    document = Document(str(source))
    if document.inline_shapes:
        return failure("docx files with embedded images are not supported")

    lines: list[str] = []
    for block in _iter_blocks(document):
        if block["type"] == "paragraph":
            text = block["text"].strip()
            if not text:
                continue
            if output_type == "markdown" and block["style"].startswith("Heading "):
                level = block["style"].split("Heading ", 1)[1]
                if level.isdigit():
                    lines.append(f"{'#' * min(int(level), 6)} {text}")
                    lines.append("")
                    continue
            lines.append(text)
            lines.append("")
        elif block["type"] == "table":
            rows = block["rows"]
            if output_type == "markdown":
                lines.extend(_markdown_table(rows))
            else:
                lines.extend("\t".join(row) for row in rows)
            lines.append("")

    rendered = "\n".join(lines).strip() + "\n"
    if not rendered.strip():
        return failure("docx contains no supported text content")
    output.write_text(rendered, encoding="utf-8")
    return success(str(output), ["pandoc unavailable; used python-docx"])


def _iter_blocks(document) -> list[dict]:
    blocks: list[dict] = []
    for paragraph in document.paragraphs:
        blocks.append({"type": "paragraph", "text": paragraph.text, "style": paragraph.style.name})
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        blocks.append({"type": "table", "rows": rows})
    return blocks


def _markdown_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    result = ["| " + " | ".join(header) + " |"]
    result.append("| " + " | ".join(["---"] * width) + " |")
    for row in normalized[1:]:
        result.append("| " + " | ".join(row) + " |")
    return result
