from __future__ import annotations

from pathlib import Path

from ..result import failure, success
from .pandoc import try_pandoc


def convert(source: Path, output: Path, input_type: str, output_type: str):
    if try_pandoc(source, output, input_type, output_type):
        return success(str(output), ["converted with pandoc"])

    from bs4 import BeautifulSoup
    from docx import Document

    document = Document()
    text = source.read_text(encoding="utf-8")

    if input_type == "txt":
        for block in _paragraph_blocks(text):
            document.add_paragraph(block)
    elif input_type == "markdown":
        _markdown_to_docx(document, text)
    elif input_type == "html":
        soup = BeautifulSoup(text, "html.parser")
        if soup.find(["script", "style", "svg", "canvas", "iframe"]):
            return failure("complex html is not supported for docx conversion")
        _html_to_docx(document, soup)
    else:
        return failure(f"unsupported docx source: {input_type}")

    document.save(str(output))
    return success(str(output), ["pandoc unavailable; used python-docx"])


def _paragraph_blocks(text: str) -> list[str]:
    return [block.strip() for block in text.split("\n\n") if block.strip()]


def _markdown_to_docx(document, text: str) -> None:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 6)
            document.add_heading(line.lstrip("#").strip(), level=level)
            index += 1
            continue
        if _is_table_start(lines, index):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(set(cell) <= {"-", ":"} and cell for cell in row):
                    rows.append(row)
                index += 1
            _add_table(document, rows)
            continue
        document.add_paragraph(line)
        index += 1


def _is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and set(lines[index + 1].strip().replace("|", "").replace(" ", "")) <= {"-", ":"}
    )


def _html_to_docx(document, soup) -> None:
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"], recursive=True):
        if element.find_parent(["table"]) and element.name != "table":
            continue
        if element.name.startswith("h"):
            document.add_heading(element.get_text(" ", strip=True), level=int(element.name[1]))
        elif element.name == "li":
            document.add_paragraph(element.get_text(" ", strip=True), style="List Bullet")
        elif element.name == "table":
            rows = [[cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])] for row in element.find_all("tr")]
            _add_table(document, rows)
        else:
            text = element.get_text(" ", strip=True)
            if text:
                document.add_paragraph(text)


def _add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
